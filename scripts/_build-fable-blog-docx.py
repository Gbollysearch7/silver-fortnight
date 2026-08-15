from docx import Document
from docx.shared import Pt, RGBColor, Inches

doc = Document()
for sec in doc.sections:
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)
nm = doc.styles['Normal']; nm.font.name = 'Calibri'; nm.font.size = Pt(11); nm.font.color.rgb = RGBColor(0,0,0)

def H(t, size=15, sb=14):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0,0,0); return p

def para(runs, sa=8):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa)
    for txt, bold in runs:
        r = p.add_run(txt); r.bold = bold; r.font.color.rgb = RGBColor(0,0,0)
    return p

def quote(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.space_after = Pt(8)
    r = p.add_run('"' + t + '"'); r.italic = True; r.font.color.rgb = RGBColor(70,70,70); return p

def bullet(t):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.color.rgb = RGBColor(0,0,0); return p

import re
md = open('how-fable-5-fixed-my-technical-seo.md').read()
lines = md.split('\n')
i = 0
while i < len(lines):
    ln = lines[i]
    if ln.startswith('# '):
        p = doc.add_paragraph(); r = p.add_run(ln[2:]); r.bold = True; r.font.size = Pt(20)
        p.paragraph_format.space_after = Pt(4)
    elif ln.startswith('## '):
        H(ln[3:])
    elif ln.startswith('> '):
        quote(ln[2:].strip('"'))
    elif ln.startswith('- '):
        bullet(ln[2:])
    elif ln.strip():
        # bold-lead paragraphs like **X.** rest
        segs = re.split(r'(\*\*.*?\*\*)', ln)
        runs = []
        for s in segs:
            if s.startswith('**') and s.endswith('**'): runs.append((s[2:-2], True))
            elif s: runs.append((s, False))
        if runs: para(runs)
    i += 1

doc.save('How-Fable-5-Fixed-My-Technical-SEO.docx')
print('Saved docx')
