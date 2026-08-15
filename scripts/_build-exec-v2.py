from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
nm = doc.styles['Normal']; nm.font.name='Calibri'; nm.font.size=Pt(11); nm.font.color.rgb=RGBColor(0,0,0)

def H(t,size=14,sb=12):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(size); r.underline=False; r.font.color.rgb=RGBColor(0,0,0); return p
def line(t,italic=False,size=11,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(t); r.italic=italic; r.font.size=Pt(size); r.underline=False
    r.font.color.rgb=RGBColor(*(color or (0,0,0))); return p
def B(t,lead=None):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3)
    if lead:
        r=p.add_run(lead); r.bold=True; r.underline=False; r.font.color.rgb=RGBColor(0,0,0)
        r2=p.add_run(t); r2.underline=False; r2.font.color.rgb=RGBColor(0,0,0)
    else:
        r=p.add_run(t); r.underline=False; r.font.color.rgb=RGBColor(0,0,0)
    return p
def table(headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        t.rows[0].cells[i].paragraphs[0].clear()
        r=t.rows[0].cells[i].paragraphs[0].add_run(h); r.bold=True; r.underline=False; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0,0,0)
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            c[i].paragraphs[0].clear()
            r=c[i].paragraphs[0].add_run(str(v)); r.underline=False; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0,0,0)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

# title
p=doc.add_paragraph(); r=p.add_run('SEO Plan and Execution'); r.bold=True; r.font.size=Pt(18); r.underline=False; r.font.color.rgb=RGBColor(0,0,0)
line('Where the blog is now, what I am keeping an eye on, and the plan to grow it.', italic=True)
line('Owner: me  ·  3 posts a week  ·  June 2026', color=(90,90,90), size=9.5)

# 1. WHERE WE ARE
H('Where we are now')
line('The blog has grown a lot in the last six months. These are blog pages only, and I have taken out branded searches (people typing "TradersYard"), so this is the traffic we are earning purely from the content.')
B('went from about 1,730 a month to over 12,000. That is roughly six times more visibility.', lead='Impressions: ')
B('we now show up for over 1,000 different searches, up from about 160.', lead='Queries we rank for: ')
B('we moved from barely showing up to ranking on the first page for a lot of terms.', lead='Positions: ')
B('clicks are still very low. We are being seen but not clicked yet. That is the main thing to fix next, and it is fixable.', lead='Clicks: ')
table(['Blog only, non-branded','Six months ago','Now'],
      [['Impressions per month','1,730','12,057'],
       ['Searches we rank for','164','1,022'],
       ['Clicks per month','0','13'],
       ['Pages that rank but get few clicks','few','around 58']])

# 2. WHY CLICKS ARE LOW
H('Why the clicks are low (and why that is good news)')
line('A lot of our pages already rank on the first page of Google. The problem is the title and the short description that show up in the search results are not pulling people in. So we get seen but not clicked.')
B('the ranking already exists, so we do not have to start from scratch.', lead='The win is quick because ')
B('we just rewrite the titles and descriptions to match what people are actually looking for.', lead='The fix is simple: ')
B('this is the fastest way to turn the visibility we have built into real clicks.', lead='The payoff: ')
line('Examples of pages that rank well but are barely clicked:')
table(['Page','Position','Seen (impressions)','Clicked'],
      [['Which prop firm gives a real account','8','5,884','8'],
       ['Prop firm copy trading','8','4,312','2'],
       ['How many people get payouts','8','4,073','21'],
       ['Best prop firms for day trading','9','2,508','3']])

# 3. WHAT I KEEP MONITORING
H('What I keep an eye on to keep it improving')
B('to see which pages are going up or down and catch new chances early.', lead='I check Search Console every week ')
B('and rewrite the titles on the pages that rank well but are not clicked.', lead='I look for click gaps ')
B('on page two and push them onto page one with better content and internal links.', lead='I find pages stuck ')
B('I have already updated, because some still need another pass to fully answer what people search.', lead='I keep improving posts ')
B('against the targets, and adjust the plan based on what is actually working.', lead='I review every month ')

# 4. THE PLAN TO GET THERE
H('The plan to get there')
line('The work splits into three parts: fix what we already have, keep publishing, and grow into new topics.')
line('Fix what we have (first, fastest wins):', italic=True)
B('rewrite the titles and descriptions on the ~58 pages that rank but are not clicked.')
B('push our page-two pages onto page one.')
B('go back over posts I improved and take them further where needed.')
line('Keep publishing (steady, every week):', italic=True)
B('publish 3 new posts a week, written around keywords we can realistically rank for.')
B('add images to articles to make them easier to read. (This is on the list, not done yet.)')
B('link articles to each other so ranking strength flows to the pages that matter.')
line('Grow into new topics (the bigger picture):', italic=True)
B('work through a researched list of 676 new pages, easiest keywords first.')
B('build out the remaining topic areas (firm reviews, platforms, pricing).')
B('get the old duplicate blog address redirected (needs Cloudflare access).')

# 5. TASKS AND TIMELINES
H('Tasks and timelines')
table(['Task','When'],
      [['Set up weekly Search Console check','Week 1'],
       ['Rewrite titles on the ~58 low-click pages','Weeks 1 to 6'],
       ['Push page-two pages onto page one','Weeks 2 to 10'],
       ['Publish 3 posts a week','Ongoing'],
       ['Keep improving older posts','Ongoing'],
       ['Add images to articles','Ongoing (not started)'],
       ['Build remaining topic areas','Months 1 to 3'],
       ['Redirect old blog address','When Cloudflare access is given'],
       ['Monthly review against targets','Every month']])

# 6. TARGETS
H('Targets (blog only, non-branded)')
line('I will measure against where we are now. These grow over time as new content ranks and the title fixes land.')
table(['Target','Now','3 months','6 months','12 months'],
      [['Impressions per month','12,000','20,000','35,000','70,000'],
       ['Clicks per month','13','250','700','2,000'],
       ['Searches we rank for','1,022','1,400','1,900','3,000'],
       ['New posts published','0','36','72','144']])
line('The full keyword list and posting schedule are in the spreadsheet (SEO Keyword Plan).', italic=True)

doc.save('TradersYard-SEO-Plan-and-Execution.docx')
print('Saved')
