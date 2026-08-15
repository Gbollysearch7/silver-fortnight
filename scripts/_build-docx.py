from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# base styles: simple, black text, no underline anywhere
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0,0,0)

def H(text, size=15, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0,0,0)
    r.underline = False
    return p

def P(text, size=11, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = italic; r.bold = bold
    r.font.size = Pt(size); r.underline = False
    r.font.color.rgb = RGBColor(0,0,0)
    return p

def B(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; r.underline = False; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0,0,0)
        r2 = p.add_run(text); r2.underline=False; r2.font.size=Pt(11); r2.font.color.rgb=RGBColor(0,0,0)
    else:
        r = p.add_run(text); r.underline=False; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0,0,0)
    return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'   # plain grid, no color
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].paragraphs[0].clear()
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True; r.underline=False; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0,0,0)
    for row in rows:
        cells = t.add_row().cells
        for i,val in enumerate(row):
            cells[i].paragraphs[0].clear()
            r = cells[i].paragraphs[0].add_run(str(val))
            r.underline=False; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0,0,0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

# ---------- TITLE ----------
title = doc.add_paragraph()
tr = title.add_run('TradersYard Blog SEO: Work in Progress')
tr.bold = True; tr.font.size = Pt(20); tr.underline=False; tr.font.color.rgb=RGBColor(0,0,0)
sub = doc.add_paragraph()
sr = sub.add_run('Current status, the data I am working from, and the tasks in progress')
sr.italic = True; sr.font.size = Pt(11); sr.underline=False; sr.font.color.rgb=RGBColor(0,0,0)
dl = doc.add_paragraph()
dr = dl.add_run('Prepared June 2026  |  Source: Google Search Console')
dr.font.size = Pt(9.5); dr.underline=False; dr.font.color.rgb=RGBColor(90,90,90)

# ---------- 1. WHERE THINGS STAND ----------
H('1. Where the blog stands today')
P('The blog has built strong search visibility over the last six months. The footprint has grown significantly, and we now appear in Google for a large number of queries. The current focus is converting that visibility into actual clicks and traffic.')
table(['Metric (28-day window)', 'Six months ago', 'Now'],
      [['Search impressions', '10,780', '55,922'],
       ['Distinct queries ranking', '386', '1,209'],
       ['Blog pages ranking', '47', '210'],
       ['Average position', '19.5', '12.1'],
       ['Clicks', '2,261', '1,559']])
P('Note on clicks: visibility (impressions) has grown roughly five-fold while clicks have not yet followed. This is expected during a build phase. We have captured a much larger audience on Google; the next stage of work is turning that visibility into clicks. The tasks below are how we do that.')

# ---------- 2. THE BIGGEST OPPORTUNITY ----------
H('2. The biggest opportunity right now: titles and headers')
P('Many of our pages already rank on the first page of Google but receive almost no clicks, because the page title and header shown in the search results are not compelling enough. Fixing the title and meta description on these pages is the fastest available win: the ranking already exists, we just need to earn the click.')
P('Examples of pages that rank well but are not being clicked (these are the priority for title and header rewrites):', bold=True)
table(['Page', 'Avg position', 'Impressions / 28d', 'Clicks / 28d'],
      [['Which prop firm gives a real account', '8.2', '5,884', '8'],
       ['Prop firm copy trading', '8.2', '4,312', '2'],
       ['Prop firm demo account and practice', '10.1', '4,170', '21'],
       ['How many people get payouts from prop firms', '7.7', '4,073', '21'],
       ['Best prop firms for day trading 2026', '9.1', '2,508', '3'],
       ['Prop firm refund policy comparison', '6.1', '2,339', '8'],
       ['Are prop firm fees tax deductible', '8.1', '877', '2'],
       ['Prop firms in Europe', '10.1', '792', '6']])
P('Across the blog there are roughly 58 pages in this position: ranking on or near page one, but with a very low click rate. Rewriting their titles and headers to match what searchers actually want is a clear, measurable task list.')

# ---------- 3. STRIKING DISTANCE ----------
H('3. Pages one step from page one')
P('A second group of pages rank on page two (positions 11 to 20). These collect impressions but almost no clicks, because few people scroll to page two. A focused push (stronger on-page content and internal links) can move them onto page one, where the clicks are.')
table(['Query', 'Avg position', 'Impressions / 28d'],
      [['Prop firm real capital', '12.6', '196'],
       ['Day trading prop firms', '14.3', '90'],
       ['Most trusted prop firms in Europe', '14.3', '89'],
       ['Demo funded account', '12.0', '76'],
       ['Prop firm practice account', '10.8', '67'],
       ['Funded account demo', '12.8', '52']])

# ---------- 4. TASKS IN PROGRESS ----------
H('4. SEO tasks in progress')
P('The current and upcoming work, in priority order:', bold=True)
B('rewrite the titles and meta descriptions on the ~58 pages that rank well but are not being clicked, so the search result earns the click.', bold_lead='Title and header rewrites — ')
B('strengthen the page-two pages so they move onto page one for queries we already rank for.', bold_lead='Push striking-distance pages — ')
B('continue improving the depth and quality of posts. Some posts I have already improved still need a further pass to fully satisfy the search intent and rank higher.', bold_lead='Improve existing content — ')
B('add supporting images and visuals to articles to improve readability and time on page. This is not yet done and is on the to-do list.', bold_lead='In-article images — ')
B('keep internal links flowing from supporting articles up to the main money pages, so ranking strength is passed where it matters.', bold_lead='Internal linking — ')
B('begin publishing new pages targeting low-competition, high-intent keywords from a planned keyword pipeline, building topical authority.', bold_lead='Expand keyword coverage — ')
B('an old duplicate blog address still needs a redirect at the infrastructure level (requires Cloudflare access) to consolidate ranking strength. Flagged and pending access.', bold_lead='Technical clean-up — ')

# ---------- 5. WHY THIS APPROACH ----------
H('5. Why this is the right sequence')
P('Search rankings are an asset that compounds. We have spent the recent phase building visibility and raising content quality across the blog. With that foundation in place, the next phase is about capturing the clicks from the visibility we already have, then expanding into new keywords. The title and header rewrites in section 2 are the quickest measurable win and are the immediate priority.')

doc.save('TradersYard-SEO-Process.docx')
print('Saved TradersYard-SEO-Process.docx')
