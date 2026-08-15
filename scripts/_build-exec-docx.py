from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
n = doc.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11); n.font.color.rgb=RGBColor(0,0,0)

def H(t,size=15,sb=12):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.bold=True; r.font.size=Pt(size); r.underline=False; r.font.color.rgb=RGBColor(0,0,0); return p
def P(t,italic=False,bold=False,size=11):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(t); r.italic=italic; r.bold=bold; r.font.size=Pt(size); r.underline=False; r.font.color.rgb=RGBColor(0,0,0); return p
def B(t,lead=None):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3)
    if lead:
        r=p.add_run(lead); r.bold=True; r.underline=False; r.font.color.rgb=RGBColor(0,0,0)
        r2=p.add_run(t); r2.underline=False; r2.font.color.rgb=RGBColor(0,0,0)
    else:
        r=p.add_run(t); r.underline=False; r.font.color.rgb=RGBColor(0,0,0)
    return p
def table(headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        t.rows[0].cells[i].paragraphs[0].clear()
        r=t.rows[0].cells[i].paragraphs[0].add_run(h); r.bold=True; r.underline=False; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0,0,0)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].paragraphs[0].clear()
            r=cells[i].paragraphs[0].add_run(str(v)); r.underline=False; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0,0,0)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return t

# TITLE
p=doc.add_paragraph(); r=p.add_run('Roadmap Execution Plan — Section 1: SEO & Organic Channel'); r.bold=True; r.font.size=Pt(19); r.underline=False; r.font.color.rgb=RGBColor(0,0,0)
p=doc.add_paragraph(); r=p.add_run('How we grow the organic search channel: current state, what we monitor, and the tasks, owner and timelines to execute.'); r.italic=True; r.font.size=Pt(11); r.underline=False; r.font.color.rgb=RGBColor(0,0,0)
p=doc.add_paragraph(); r=p.add_run('Owner: SEO Lead   |   Cadence: 3 posts per week   |   Prepared June 2026'); r.font.size=Pt(9.5); r.underline=False; r.font.color.rgb=RGBColor(90,90,90)

# 1. CURRENT STATE (simplified)
H('1. Current state (simplified)')
P('Over the last six months the blog has grown its search visibility substantially. We now appear in Google for over 1,200 search queries across 210 pages. The challenge is converting that visibility into clicks, because many pages rank well but are not yet being clicked.')
table(['Where we are','Figure'],
      [['Search impressions per month','55,900'],
       ['Blog pages ranking in Google','210'],
       ['Distinct queries ranking','1,209'],
       ['Average position','12.1 (was 19.5 six months ago)'],
       ['Pages that rank well but get few clicks','~58 (the immediate opportunity)']])
P('In one line: strong visibility built, now we capture the clicks and keep expanding.')

# 2. WHAT WE DO TO MONITOR & IMPROVE
H('2. How we monitor and improve this stage')
P('This channel is run on a continuous loop, not a one-off project. The monitoring keeps it improving every month:', bold=True)
B('every week, using Google Search Console, to spot pages that gained or lost rankings and catch new opportunities early.', lead='Track rankings — ')
B('the pages that rank on page one but have a low click rate, and rewrite their titles and descriptions to earn the click.', lead='Find the click gaps — ')
B('the pages sitting on page two (positions 11-20) and give them a push with stronger content and internal links to move them onto page one.', lead='Push striking-distance pages — ')
B('content quality on a rolling basis. Some posts have been improved already but still need a further pass to fully satisfy search intent.', lead='Keep improving existing posts — ')
B('a monthly performance review against the KPI targets, adjusting the plan based on what is actually ranking.', lead='Review monthly — ')

# 3. EXECUTION TASKS
H('3. Execution plan — tasks, owner and timelines')
P('The work is organised into immediate, ongoing, and growth tasks. All owned by the SEO Lead.', bold=True)
table(['Task','Owner','Timeline','Priority'],
      [['Rewrite titles and headers on the ~58 pages that rank well but are not clicked','SEO Lead','Weeks 1-6','High — fastest win'],
       ['Set up weekly Search Console tracking routine','SEO Lead','Week 1','High'],
       ['Push page-two pages onto page one (content + internal links)','SEO Lead','Weeks 2-10','High'],
       ['Further-improve already-updated posts (rolling pass)','SEO Lead','Ongoing','Medium'],
       ['Add supporting images to articles','SEO Lead','Ongoing (not yet started)','Medium'],
       ['Publish 3 new keyword-targeted posts per week','SEO Lead','Ongoing','High'],
       ['Build out remaining content pillars (Firm Reviews, Platforms, Pricing)','SEO Lead','Months 1-3','Medium'],
       ['Old duplicate blog address redirect','SEO Lead (needs Cloudflare access)','When access granted','Medium'],
       ['Monthly KPI review and plan adjustment','SEO Lead','Monthly','High']])

# 4. KEYWORD PIPELINE
H('4. The keyword pipeline (where the new posts come from)')
P('We have a researched pipeline of 676 net-new pages to publish, targeting 58,330 searches per month combined. They are ordered easiest-first so we build authority before chasing harder terms. The full list, posting schedule, and KPI targets are in the attached spreadsheet (TradersYard-SEO-Keyword-Plan.xlsx).')
table(['Wave','Difficulty','Pages','Monthly volume'],
      [['Wave 1 — Quick Wins','Very low (KD 0-10)','318','18,520'],
       ['Wave 2 — Easy','Low (KD 11-20)','210','18,560'],
       ['Wave 3 — Medium','Moderate (KD 21-30)','148','21,250']])

# 5. KPIs
H('5. KPI targets')
P('Success is measured against the current Search Console baseline. Targets grow over time as new content ranks and compounds.')
table(['KPI','Now','3 months','6 months','12 months'],
      [['Impressions / month','55,900','70,000','95,000','160,000'],
       ['Blog clicks / month','1,560','2,200','3,500','7,000'],
       ['Queries ranking','1,209','1,500','2,000','3,200'],
       ['Pages ranking','210','250','320','450'],
       ['Average position','12.1','11.0','9.5','8.0'],
       ['New posts (cumulative)','0','36','72','144']])

doc.save('TradersYard-SEO-Execution-Plan.docx')
print('Saved TradersYard-SEO-Execution-Plan.docx')
